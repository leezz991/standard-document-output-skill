#!/usr/bin/env python3
"""Build a standard DOCX in official or general mode from simple Markdown."""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


STYLE_BY_LEVEL = {1: "Title", 2: "Heading 1", 3: "Heading 2", 4: "Heading 3", 5: "Heading 4"}
IMAGE_FILENAME_RE = re.compile(r"^ChatGPT Image .*\.(?:png|jpe?g|gif|webp)$", re.IGNORECASE)
IMAGE_MARKDOWN_RE = re.compile(r"^!\[([^]]*)\]\(([^)]+)\)$")


def is_table_row(text: str) -> bool:
    return text.startswith("|") and text.endswith("|") and text.count("|") >= 2


def table_cells(text: str) -> list[str]:
    return [cell.strip() for cell in text.strip()[1:-1].split("|")]


def is_table_separator(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def parse_markdown(path: Path, mode: str) -> list[tuple[str, object]]:
    items: list[tuple[str, object]] = []
    lines = path.read_text(encoding="utf-8-sig").splitlines()
    index = 0
    while index < len(lines):
        text = lines[index].strip()
        if is_table_row(text) and index + 1 < len(lines) and is_table_row(lines[index + 1].strip()):
            header, separator = table_cells(text), table_cells(lines[index + 1].strip())
            if is_table_separator(separator):
                rows = [header]
                index += 2
                while index < len(lines) and is_table_row(lines[index].strip()):
                    row = table_cells(lines[index].strip())
                    if len(row) != len(header):
                        raise ValueError("Every Markdown table row must have the same number of columns.")
                    rows.append(row)
                    index += 1
                items.append(("__table__", rows))
                continue
        if not text or IMAGE_FILENAME_RE.match(text):
            index += 1
            continue
        image_match = IMAGE_MARKDOWN_RE.match(text)
        if image_match:
            image_path = Path(image_match.group(2)).expanduser()
            if not image_path.is_absolute():
                image_path = path.parent / image_path
            if not image_path.is_file():
                raise ValueError(f"Image file not found: {image_path}")
            items.append(("__image__", image_path))
            index += 1
            continue
        if text.startswith("> "):
            if mode != "general":
                raise ValueError("Subtitle marker '> ' is only supported in general mode.")
            items.append(("图表内容", text[2:].strip()))
            index += 1
            continue
        level = len(text) - len(text.lstrip("#"))
        items.append((STYLE_BY_LEVEL[level], text[level:].strip()) if 1 <= level <= 5 and len(text) > level and text[level].isspace() else ("Normal", text))
        index += 1
    if not items:
        raise ValueError("Input contains no non-empty paragraphs.")
    return items


def apply_general_format(paragraph, style_name: str) -> None:
    fmt = paragraph.paragraph_format
    if style_name == "Title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before, fmt.space_after = Pt(5), Pt(5)
    elif style_name == "Heading 1":
        fmt.space_before, fmt.space_after = Pt(15.6), Pt(7.8)
    elif style_name == "Normal":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent, fmt.line_spacing = Pt(24), 1.5


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")


def add_general_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment, table.autofit = WD_TABLE_ALIGNMENT.CENTER, False
    weights = [max(4, max(len(row[column]) for row in rows)) for column in range(len(rows[0]))]
    for row_index, row in enumerate(rows):
        for column, text in enumerate(row):
            cell = table.cell(row_index, column)
            cell.width = Cm(16.0 * weights[column] / sum(weights))
            set_cell_borders(cell)
            paragraph = cell.paragraphs[0]
            paragraph.text, paragraph.style = text, "图表内容"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if column == 0 or max(len(item[column]) for item in rows) <= 12 else WD_ALIGN_PARAGRAPH.LEFT


def add_general_image(document: Document, image_path: Path) -> None:
    with Image.open(image_path) as image:
        dpi_x = image.info.get("dpi", (96, 96))[0] or 96
        natural_width_cm = image.width / dpi_x * 2.54
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.98)) if natural_width_cm > 15.98 else run.add_picture(str(image_path))


def clear_body(document: Document) -> None:
    for child in list(document._element.body):
        if not child.tag.endswith("}sectPr"):
            document._element.body.remove(child)


def build(input_path: Path, output_path: Path, template_path: Path, mode: str) -> None:
    if input_path.resolve() == output_path.resolve():
        raise ValueError("Input and output paths must differ.")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    document = Document(output_path)
    clear_body(document)
    for style_name, content in parse_markdown(input_path, mode):
        if style_name == "__table__":
            if mode != "general":
                raise ValueError("Markdown tables are currently supported only in general mode.")
            add_general_table(document, content)
        elif style_name == "__image__":
            add_general_image(document, content)
        else:
            paragraph = document.add_paragraph(content, style=style_name)
            if mode == "general":
                apply_general_format(paragraph, style_name)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="UTF-8 Markdown source")
    parser.add_argument("output", type=Path, help="Output .docx path")
    parser.add_argument("--mode", required=True, choices=("official", "general"))
    parser.add_argument("--template", type=Path, help="Optional template override")
    args = parser.parse_args()
    if args.output.suffix.lower() != ".docx":
        parser.error("output must use the .docx extension")
    template = args.template or (Path(__file__).resolve().parents[1] / "assets" / ("gongwen.docx" if args.mode == "official" else "reference.docx"))
    build(args.input, args.output, template, args.mode)


if __name__ == "__main__":
    main()
